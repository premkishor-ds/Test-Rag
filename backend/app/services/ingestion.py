import os
import re
import csv
import json
import hashlib
import datetime
import logging
import urllib.parse
from urllib.parse import urlparse, parse_qs
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session

# Import third-party libraries for parsing
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

from qdrant_client.http import models as qmodels
from app.core.config import settings
from app.core.ollama import ollama_client
from app.core.qdrant import qdrant_client
from app.models.models import Stock, AnnualReport, CorporateDocument, AuditLog

import requests

# FIX: logger must be defined BEFORE any function that uses it
logger = logging.getLogger(__name__)

def fetch_stock_data_from_yahoo(symbol: str) -> dict:
    yf_symbol = symbol.upper()
    if "." not in yf_symbol:
        yf_symbol = f"{yf_symbol}.NS"
        
    try:
        import yfinance as yf
        logger.info(f"Fetching data for {yf_symbol} via yfinance...")
        ticker = yf.Ticker(yf_symbol)
        info = ticker.info
        if not info or not info.get("longName"):
            logger.warning(f"No info returned for {yf_symbol}")
            return {}
            
        name = info.get("longName") or info.get("shortName") or symbol
        sector = info.get("sector", "Unknown")
        industry = info.get("industry", "Unknown")
        
        raw_mcap = info.get("marketCap", 0.0)
        market_cap = round(raw_mcap / 10000000.0, 2) if raw_mcap else 0.0
        
        roe = round((info.get("returnOnEquity") or 0.0) * 100, 2)
        roce = round((info.get("returnOnAssets") or 0.0) * 100 * 1.3, 2)
        
        raw_de = info.get("debtToEquity", 0.0)
        debt_to_equity = round(raw_de / 100.0, 2) if raw_de else 0.0
        
        raw_cf = info.get("operatingCashflow", 0.0)
        cash_flow = round(raw_cf / 10000000.0, 2) if raw_cf else 0.0
        
        rev_growth = round((info.get("revenueGrowth") or 0.0) * 100, 2)
        prof_growth = round((info.get("earningsGrowth") or 0.0) * 100, 2)
        if prof_growth == 0.0:
            prof_growth = round(rev_growth * 1.1, 2)
            
        pe_ratio = info.get("forwardPE") or info.get("trailingPE")
        peg_ratio = info.get("pegRatio")
        
        raw_rev = info.get("totalRevenue", 0.0)
        revenue = round(raw_rev / 10000000.0, 2) if raw_rev else 0.0
        net_profit = round(revenue * (info.get("profitMargins", 0.05) or 0.05), 2)
        
        current_price = info.get("currentPrice") or info.get("regularMarketPrice") or 1.0
        fifty_high = info.get("fiftyTwoWeekHigh") or (current_price * 1.2)
        fifty_low = info.get("fiftyTwoWeekLow") or (current_price * 0.8)
        
        return {
            "name": name,
            "sector": sector,
            "industry": industry,
            "market_cap": market_cap,
            "roe": roe,
            "roce": roce if roce > 0.0 else roe,
            "debt_to_equity": debt_to_equity,
            "cash_flow": cash_flow,
            "revenue_growth": rev_growth,
            "profit_growth": prof_growth,
            "pe_ratio": round(pe_ratio, 2) if pe_ratio else None,
            "peg_ratio": round(peg_ratio, 2) if peg_ratio else None,
            "revenue": revenue,
            "net_profit": net_profit,
            "fifty_two_week_high": round(fifty_high, 2),
            "fifty_two_week_low": round(fifty_low, 2),
            "current_price": current_price
        }
    except Exception as e:
        logger.error(f"Error fetching data from Yahoo Finance for {symbol}: {e}")
        return {}

class DocumentExtractor:
    @staticmethod
    def extract_text(file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".txt":
            return DocumentExtractor.extract_txt(file_path)
        elif ext == ".pdf":
            return DocumentExtractor.extract_pdf(file_path)
        elif ext == ".docx":
            return DocumentExtractor.extract_docx(file_path)
        elif ext == ".html" or ext == ".htm":
            return DocumentExtractor.extract_html(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return DocumentExtractor.extract_image_ocr(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext} for path {file_path}")
            return ""

    @staticmethod
    def extract_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return ""

    @staticmethod
    def extract_pdf(file_path: str) -> str:
        if not pypdf:
            logger.error("pypdf library not installed. Cannot parse PDF.")
            return ""
        try:
            text = []
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            return ""

    @staticmethod
    def extract_docx(file_path: str) -> str:
        if not docx:
            logger.error("python-docx library not installed. Cannot parse DOCX.")
            return ""
        try:
            doc = docx.Document(file_path)
            text = [p.text for p in doc.paragraphs]
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            return ""

    @staticmethod
    def extract_html(file_path: str) -> str:
        if not BeautifulSoup:
            logger.error("beautifulsoup4 library not installed. Cannot parse HTML.")
            return ""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                return soup.get_text(separator="\n")
        except Exception as e:
            logger.error(f"Error reading HTML file {file_path}: {e}")
            return ""

    @staticmethod
    def extract_image_ocr(file_path: str) -> str:
        if not Image or not pytesseract:
            logger.error("Pillow or pytesseract libraries not installed or Tesseract missing. Cannot perform OCR.")
            return ""
        try:
            img = Image.open(file_path)
            # Simple pytesseract OCR execution
            text = pytesseract.image_to_string(img)
            return text
        except Exception as e:
            logger.error(f"Error executing OCR on image {file_path}: {e}")
            return ""

class TextChunker:
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        chunks = []
        if len(text) <= chunk_size:
            return [text] if text else []

        start = 0
        while start < len(text):
            end = start + chunk_size
            # If we are not at the end of the text, try to find a natural boundary (sentence or space)
            if end < len(text):
                # Look back up to 100 characters for a sentence end or space
                boundary = -1
                for i in range(end, max(end - 100, start + overlap), -1):
                    if text[i] in ['.', '!', '?']:
                        boundary = i + 1
                        break
                if boundary == -1:
                    # Fallback to space
                    for i in range(end, max(end - 100, start + overlap), -1):
                        if text[i] == ' ':
                            boundary = i
                            break
                if boundary != -1:
                    end = boundary
            
            chunks.append(text[start:end].strip())
            start = end - overlap
            if start < 0:
                start = 0
            # Ensure forward progress
            if end <= start:
                start = end
        return chunks

class IngestionEngine:
    def __init__(self, db: Session):
        self.db = db

    def sync_stocks_from_csv(self) -> List[Stock]:
        csv_path = settings.STOCKS_CSV
        if not os.path.exists(csv_path):
            # Create an empty stocks CSV template if it doesn't exist
            os.makedirs(os.path.dirname(csv_path), exist_ok=True)
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["symbol", "name", "sector", "industry", "market_cap"])
            logger.info(f"Created empty stocks template CSV at {csv_path}. Please populate with your stocks.")

        stocks_loaded = []
        with open(csv_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                symbol = row.get("symbol", "").strip().upper()
                if not symbol:
                    continue
                
                name = row.get("name", "").strip() if row.get("name") else ""
                
                # Check if stock exists
                stock = self.db.query(Stock).filter(Stock.symbol == symbol).first()
                if not stock:
                    if not name:
                        logger.info(f"Enriching metadata for symbol: {symbol} from Yahoo Finance API...")
                        yf_data = fetch_stock_data_from_yahoo(symbol)
                        if yf_data and yf_data.get("name") != symbol:
                            name = yf_data["name"]
                            sector = yf_data["sector"]
                            industry = yf_data["industry"]
                            mcap = yf_data["market_cap"]
                            logger.info(f"Successfully retrieved profile for {symbol} from Internet: {name}")
                        else:
                            logger.info(f"Internet query returned empty for {symbol}. Falling back to local AI...")
                            try:
                                prompt = (
                                    "You are a financial metadata enricher. Provide standard market data for the Indian stock symbol: "
                                    f"\"{symbol}\". Return ONLY a raw JSON object with the following keys, no markdown code block formatting:\n"
                                    "{\n"
                                    "  \"name\": \"Full Company Name\",\n"
                                    "  \"sector\": \"Standard Sector\",\n"
                                    "  \"industry\": \"Standard Industry\",\n"
                                    "  \"market_cap\": 50000.0\n"
                                    "}"
                                )
                                enriched_raw = ollama_client.generate_completion(prompt)
                                clean_json = enriched_raw.strip()
                                if clean_json.startswith("```json"):
                                    clean_json = clean_json[7:]
                                if clean_json.endswith("```"):
                                    clean_json = clean_json[:-3]
                                enriched_data = json.loads(clean_json.strip())
                                
                                name = enriched_data.get("name", f"Company {symbol}")
                                sector = enriched_data.get("sector", "Unknown")
                                industry = enriched_data.get("industry", "Unknown")
                                mcap = float(enriched_data.get("market_cap", 0.0))
                            except Exception as e:
                                logger.error(f"Failed to auto-enrich metadata for {symbol} via LLM fallback: {e}")
                                name = f"Company {symbol}"
                                sector = "Unknown"
                                industry = "Unknown"
                                mcap = 0.0
                    else:
                        sector = row.get("sector", "Unknown").strip()
                        industry = row.get("industry", "Unknown").strip()
                        mcap = float(row.get("market_cap", 0.0)) if row.get("market_cap") else 0.0

                    stock = Stock(
                        symbol=symbol,
                        name=name,
                        sector=sector,
                        industry=industry,
                        market_cap=mcap
                    )
                    self.db.add(stock)
                    logger.info(f"Added stock: {symbol} - {name}")
                else:
                    if name:
                        stock.name = name
                    # If sector or industry or market cap is missing, enrich it from Yahoo Finance
                    if not stock.sector or stock.sector == "Unknown" or not stock.industry or stock.industry == "Unknown" or stock.market_cap == 0.0:
                        logger.info(f"Enriching missing metadata for existing symbol: {symbol} via yfinance...")
                        yf_data = fetch_stock_data_from_yahoo(symbol)
                        if yf_data and yf_data.get("name") != symbol:
                            if not stock.sector or stock.sector == "Unknown":
                                stock.sector = yf_data.get("sector", "Unknown")
                            if not stock.industry or stock.industry == "Unknown":
                                stock.industry = yf_data.get("industry", "Unknown")
                            if stock.market_cap == 0.0:
                                stock.market_cap = yf_data.get("market_cap", 0.0)
                    if row.get("sector"):
                        stock.sector = row.get("sector").strip()
                    if row.get("industry"):
                        stock.industry = row.get("industry").strip()
                    if row.get("market_cap"):
                        stock.market_cap = float(row.get("market_cap"))
                stocks_loaded.append(stock)
        
        # Cleanup: Delete stocks from DB that are not in the CSV
        loaded_symbols = {s.symbol for s in stocks_loaded}
        all_db_stocks = self.db.query(Stock).all()
        for db_stock in all_db_stocks:
            if db_stock.symbol not in loaded_symbols:
                logger.info(f"Removing deleted stock {db_stock.symbol} from database...")
                self.db.delete(db_stock)
                
        self.db.commit()
        return stocks_loaded

    def ingest_document(
        self,
        file_path: str,
        stock_symbol: str,
        source_type: str,
        financial_year: int,
        quarter: Optional[str] = None
    ) -> bool:
        if not os.path.exists(file_path):
            logger.error(f"Document file does not exist: {file_path}")
            return False

        stock = self.db.query(Stock).filter(Stock.symbol == stock_symbol.upper()).first()
        if not stock:
            logger.error(f"Cannot ingest document: Stock {stock_symbol} does not exist in DB.")
            return False

        # Extract text
        text = DocumentExtractor.extract_text(file_path)
        if not text:
            logger.warning(f"No content extracted from {file_path}")
            return False

        # Handle versioning in Postgres/SQLite using CorporateDocument
        existing_docs = self.db.query(CorporateDocument).filter(
            CorporateDocument.stock_symbol == stock.symbol,
            CorporateDocument.document_type == source_type,
            CorporateDocument.financial_year == financial_year,
            CorporateDocument.quarter == quarter
        ).order_by(CorporateDocument.version.desc()).all()

        version = 1
        if existing_docs:
            for d in existing_docs:
                d.is_latest = False
            version = existing_docs[0].version + 1

        new_doc = CorporateDocument(
            stock_symbol=stock.symbol,
            document_type=source_type,
            file_path=file_path,
            financial_year=financial_year,
            quarter=quarter,
            version=version,
            is_latest=True,
            summary=f"Ingested from {os.path.basename(file_path)}"
        )
        self.db.add(new_doc)
        
        # Keep writing to AnnualReport if source_type is annual_report to maintain backward compatibility
        if source_type == "annual_report":
            existing_reports = self.db.query(AnnualReport).filter(
                AnnualReport.stock_symbol == stock.symbol,
                AnnualReport.financial_year == financial_year
            ).order_by(AnnualReport.version.desc()).all()
            for rep in existing_reports:
                rep.is_latest = False
            new_report = AnnualReport(
                stock_symbol=stock.symbol,
                file_path=file_path,
                financial_year=financial_year,
                version=version,
                is_latest=True,
                summary=f"Ingested from {os.path.basename(file_path)}"
            )
            self.db.add(new_report)

        self.db.commit()

        # Chunk document
        chunks = TextChunker.chunk_text(text, chunk_size=1000, overlap=200)
        logger.info(f"Split {file_path} into {len(chunks)} chunks.")

        # Upload vectors to Qdrant
        points = []
        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{stock.symbol}_FY{financial_year}_{quarter or 'AR'}_{version}_{i}"
            
            # Generate embedding using Ollama
            try:
                vector = ollama_client.generate_embeddings(chunk_text)
            except Exception as e:
                logger.error(f"Error vectorizing chunk {i}: {e}")
                continue

            # Construct Qdrant payload
            payload = {
                "stock_symbol": stock.symbol,
                "stock_name": stock.name,
                "source_type": source_type,
                "source_file": os.path.basename(file_path),
                "financial_year": financial_year,
                "quarter": quarter or "FY",
                "document_date": datetime.date.today().isoformat(),
                "chunk_id": chunk_id,
                "page_number": (i // 4) + 1,  # Rough page estimation if raw text
                "created_at": datetime.datetime.utcnow().isoformat(),
                "content": chunk_text
            }

            points.append(
                qmodels.PointStruct(
                    id=hashlib.md5(chunk_id.encode()).hexdigest()[:32],  # Convert to a 32-char hex UUID format
                    vector=vector,
                    payload=payload
                )
            )

        if points and qdrant_client:
            try:
                qdrant_client.upsert(
                    collection_name=settings.QDRANT_COLLECTION_NAME,
                    points=points
                )
                logger.info(f"Successfully upserted {len(points)} chunks into Qdrant collection '{settings.QDRANT_COLLECTION_NAME}'")
            except Exception as e:
                logger.error(f"Failed to upsert points to Qdrant: {e}")
                return False

        # Log audit action
        audit = AuditLog(
            action="INGEST_DOCUMENT",
            target_type=source_type,
            target_id=str(new_doc.id),
            details=f"Ingested {os.path.basename(file_path)} for stock {stock.symbol} ({source_type}, FY{financial_year}, Q{quarter or 'FY'}, Version {version})"
        )
        self.db.add(audit)
        self.db.commit()
        return True

def _try_bse_nse_direct_download(symbol: str, financial_year: int) -> List[str]:
    """
    Attempt to find PDF links from BSE India's known filing pages and
    company investor-relations patterns before falling back to web scraping.
    Returns a list of candidate PDF URLs to try.
    """
    candidate_urls = []
    # BSE annual report pattern (some companies)
    bse_url = (
        f"https://www.bseindia.com/bseplus/AnnualReport/"
        f"{symbol}/{financial_year}/AnnualReport.pdf"
    )
    candidate_urls.append(bse_url)
    # Common company IR page PDF patterns
    common_patterns = [
        f"https://www.{symbol.lower()}.com/investor-relations/annual-report-{financial_year}.pdf",
        f"https://www.{symbol.lower()}.in/annual-report-{financial_year}.pdf",
    ]
    candidate_urls.extend(common_patterns)
    return candidate_urls


def search_internet_for_pdf_links(query: str) -> List[str]:
    """
    Search Bing Search (primary), DuckDuckGo (fallback), and Yahoo Search (fallback) for PDF links.
    Returns a deduplicated list of candidate PDF URLs.
    """
    pdf_urls = []

    # FIX: URL-encode the query so spaces and special chars are handled correctly
    encoded_query = urllib.parse.quote_plus(query)

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "DNT": "1",
    }

    # 1. Try Bing Search first
    bing_url = f"https://www.bing.com/search?q={encoded_query}"
    try:
        response = requests.get(bing_url, headers=headers, timeout=15)
        if response.status_code == 200 and BeautifulSoup:
            soup = BeautifulSoup(response.text, "html.parser")
            import base64
            for a in soup.find_all("a"):
                href = a.get("href", "")
                if "bing.com/ck/a?!" in href and "&u=" in href:
                    parsed = urlparse(href)
                    u_params = parse_qs(parsed.query)
                    u_list = u_params.get("u")
                    if u_list:
                        u_val = u_list[0]
                        if u_val.startswith("a1"):
                            u_val = u_val[2:]
                        elif u_val.startswith("a0"):
                            u_val = u_val[2:]
                        # Pad base64
                        u_val += "=" * ((4 - len(u_val) % 4) % 4)
                        try:
                            decoded = base64.b64decode(u_val).decode("utf-8", errors="ignore")
                            if decoded.lower().startswith("http") and ".pdf" in decoded.lower():
                                pdf_urls.append(decoded)
                        except Exception:
                            pass
            logger.info(f"Bing Search returned {len(pdf_urls)} PDF candidate(s) for: {query}")
        else:
            logger.warning(f"Bing Search returned status {response.status_code}")
    except Exception as e:
        logger.warning(f"Bing Search failed or timed out: {e}")

    # 2. Fallback to DuckDuckGo HTML
    if not pdf_urls:
        logger.info("Bing returned no PDF links. Falling back to DuckDuckGo...")
        ddg_url = f"https://html.duckduckgo.com/html/?q={encoded_query}"
        try:
            response = requests.get(ddg_url, headers=headers, timeout=2)
            if response.status_code == 200 and BeautifulSoup:
                soup = BeautifulSoup(response.text, "html.parser")
                for css_class in ["result__a", "result__url"]:
                    anchors = soup.find_all("a", class_=css_class)
                    for a in anchors:
                        href = a.get("href", "")
                        if "uddg=" in href:
                            parsed = urlparse(href)
                            uddg_list = parse_qs(parsed.query).get("uddg")
                            if uddg_list:
                                href = urllib.parse.unquote(uddg_list[0])
                        if href and (".pdf" in href.lower()):
                            pdf_urls.append(href)
                logger.info(f"DuckDuckGo returned {len(pdf_urls)} PDF candidate(s)")
            else:
                logger.warning(f"DuckDuckGo returned status {response.status_code}")
        except Exception as e:
            logger.warning(f"DuckDuckGo search failed or timed out: {e}")

    # 3. Fallback to Yahoo Search
    if not pdf_urls:
        logger.info("Bing and DuckDuckGo returned no PDF links. Falling back to Yahoo Search...")
        yahoo_url = f"https://search.yahoo.com/search?p={encoded_query}"
        try:
            response = requests.get(yahoo_url, headers=headers, timeout=2)
            if response.status_code == 200 and BeautifulSoup:
                soup = BeautifulSoup(response.text, "html.parser")
                for a in soup.find_all("a"):
                    href = a.get("href", "")
                    if ".pdf" in href.lower():
                        if "RU=" in href:
                            match = re.search(r"RU=([^/]+)", href)
                            if match:
                                decoded = urllib.parse.unquote(match.group(1))
                                if ".pdf" in decoded.lower():
                                    pdf_urls.append(decoded)
                        elif href.lower().startswith("http") and ".pdf" in href.lower():
                            pdf_urls.append(href)
                logger.info(f"Yahoo Search returned {len(pdf_urls)} PDF candidate(s)")
            else:
                logger.warning(f"Yahoo Search returned status {response.status_code}")
        except Exception as e:
            logger.warning(f"Yahoo search failed or timed out: {e}")

    # Deduplicate
    seen = set()
    return [u for u in pdf_urls if not (u in seen or seen.add(u))]


def _download_pdf_from_urls(pdf_urls: List[str], save_path: str) -> Optional[str]:
    """
    Shared PDF download helper.
    Tries each URL in order, saves the first successful PDF to save_path.
    Returns the save_path on success, or None on total failure.
    """
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        )
    }

    for i, pdf_url in enumerate(pdf_urls):
        logger.info(f"Attempting download [{i + 1}/{len(pdf_urls)}]: {pdf_url}")
        try:
            pdf_response = requests.get(
                pdf_url, headers=headers, stream=True, timeout=30, allow_redirects=True
            )
            pdf_response.raise_for_status()

            content_type = pdf_response.headers.get("content-type", "").lower()
            url_has_pdf = ".pdf" in pdf_url.lower()

            # FIX: Accept application/pdf, application/octet-stream, binary/octet-stream
            # as long as the URL contained .pdf — many BSE/NSE servers use octet-stream
            is_pdf_content = (
                "application/pdf" in content_type
                or "application/octet-stream" in content_type
                or "binary/octet-stream" in content_type
                or url_has_pdf
            )

            if not is_pdf_content:
                logger.warning(
                    f"URL did not return PDF-like content-type ({content_type}). Skipping."
                )
                continue

            # Read response and validate it starts with %PDF magic bytes
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            tmp_path = save_path + ".tmp"
            total_bytes = 0
            with open(tmp_path, "wb") as f:
                for chunk in pdf_response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
                        total_bytes += len(chunk)

            if total_bytes < 1024:
                logger.warning(f"Downloaded file is too small ({total_bytes} bytes). Likely not a real PDF.")
                os.remove(tmp_path)
                continue

            # Validate PDF magic bytes (%PDF)
            with open(tmp_path, "rb") as f:
                magic = f.read(4)
            if magic != b"%PDF":
                logger.warning(f"File does not start with %PDF magic bytes. Not a valid PDF. Skipping.")
                os.remove(tmp_path)
                continue

            # Rename temp file to final path
            if os.path.exists(save_path):
                os.remove(save_path)
            os.rename(tmp_path, save_path)
            logger.info(f"Successfully downloaded PDF ({total_bytes / 1024:.1f} KB) → {save_path}")
            return save_path

        except Exception as download_err:
            logger.warning(f"Failed downloading from {pdf_url}: {download_err}. Trying next...")
            # Cleanup any partial temp file
            tmp_path = save_path + ".tmp"
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            continue

    return None


def _get_screener_pdf_links(symbol: str, target_type: str, financial_year: int, quarter: Optional[str] = None) -> List[str]:
    """
    Scrape Screener.in company page for PDF links.
    Categorizes based on target_type:
      - 'annual_report'
      - 'concall'
      - 'presentation'
      - 'quarterly_result'
    """
    url = f"https://www.screener.in/company/{symbol.upper()}/"
    import time
    time.sleep(1.0)
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        )
    }
    candidate_urls = []
    
    if not BeautifulSoup:
        logger.warning("BeautifulSoup not available, skipping Screener.in scraping.")
        return []

    try:
        logger.info(f"Fetching Screener.in page for {symbol} to extract {target_type}...")
        res = requests.get(url, headers=headers, timeout=10)
        if res.status_code != 200:
            logger.warning(f"Screener.in request for {symbol} returned status code {res.status_code}")
            return []
            
        soup = BeautifulSoup(res.text, "html.parser")
        year_str = str(financial_year)
        short_year = str(financial_year)[2:]  # e.g. "24" for 2024
        prev_short_year = str(financial_year - 1)[2:]  # e.g. "23" for 2024
        
        for a in soup.find_all("a"):
            href = a.get("href", "")
            text = a.get_text().strip()
            if not href:
                continue
                
            href_lower = href.lower()
            text_lower = text.lower()
            
            # Must look like a PDF or BSE/NSE filing document link
            is_doc_link = (
                href_lower.endswith(".pdf")
                or ".pdf#" in href_lower
                or "annpdfopen.aspx" in href_lower
                or "xml-data/corpfiling" in href_lower
            )
            if not is_doc_link:
                continue
                
            # Classify
            if target_type == "annual_report":
                is_annual = (
                    "financial year" in text_lower 
                    or "annual report" in text_lower 
                    or "annual-report" in href_lower 
                    or "annual_reports" in href_lower 
                    or "annualreport" in href_lower
                )
                has_year = (
                    year_str in text_lower 
                    or year_str in href_lower
                    or (short_year in text_lower and f"fy{short_year}" in text_lower.replace(" ", ""))
                    or f"{prev_short_year}_{short_year}" in href_lower
                    or f"{prev_short_year}-{short_year}" in href_lower
                    or f"{financial_year-1}-{financial_year}" in href_lower
                )
                if is_annual and has_year:
                    candidate_urls.append(href)
                    
            elif target_type in ["concall", "presentation", "quarterly_result"]:
                is_match = False
                if target_type == "concall":
                    is_match = (
                        "transcript" in text_lower 
                        or "concall" in text_lower 
                        or "con-call" in text_lower 
                        or "earning call" in text_lower 
                        or "earnings call" in text_lower 
                        or "ec_transcript" in href_lower 
                        or "transcript" in href_lower
                    )
                elif target_type == "presentation":
                    is_match = (
                        "presentation" in text_lower 
                        or "ppt" in text_lower 
                        or "fact sheet" in text_lower 
                        or "factsheet" in text_lower 
                        or "investor presentation" in text_lower 
                        or "ip_" in href_lower 
                        or "presentation" in href_lower 
                        or "fact" in href_lower
                    )
                elif target_type == "quarterly_result":
                    is_match = (
                        "result" in text_lower 
                        or "financial result" in text_lower 
                        or "notes" in text_lower 
                        or "statement" in text_lower 
                        or "result" in href_lower
                    )
                
                if is_match:
                    has_year = (
                        year_str in text_lower 
                        or year_str in href_lower 
                        or f"fy{short_year}" in text_lower.replace(" ", "") 
                        or f"_{short_year}" in href_lower 
                        or f"-{short_year}" in href_lower
                    )
                    has_quarter = True
                    if quarter:
                        q_lower = quarter.lower()
                        has_quarter = q_lower in text_lower or q_lower in href_lower
                    
                    if has_year and has_quarter:
                        candidate_urls.append(href)
                        
        logger.info(f"Screener.in extraction found {len(candidate_urls)} PDF links for {symbol} {target_type} FY{financial_year}")
    except Exception as e:
        logger.error(f"Error extracting PDF links from Screener.in for {symbol}: {e}")
        
    return candidate_urls


def fetch_and_save_pdf(symbol: str, financial_year: int) -> Optional[str]:
    """
    Download the annual report PDF for a stock symbol and financial year.
    Strategy order:
      1. Scrape Screener.in for direct links (Primary)
      2. Try BSE/NSE direct known URL patterns
      3. Search web (Bing / DDG / Yahoo) (Fallback)
    """
    logger.info(f"Fetching Annual Report PDF for {symbol} FY{financial_year}...")

    doc_dir = os.path.join(settings.DATA_DIR, "documents")
    os.makedirs(doc_dir, exist_ok=True)
    file_name = f"{symbol.upper()}_{financial_year}_AnnualReport.pdf"
    save_path = os.path.join(doc_dir, file_name)

    # Shortcut: If valid PDF already exists, reuse it
    if os.path.exists(save_path) and os.path.getsize(save_path) > 10000:
        try:
            with open(save_path, "rb") as f:
                magic = f.read(4)
                if magic == b"%PDF":
                    logger.info(f"PDF already exists on disk: {save_path}. Skipping download.")
                    return save_path
        except Exception:
            pass

    # Strategy 1: Screener.in scraper (Primary)
    screener_urls = _get_screener_pdf_links(symbol, "annual_report", financial_year)
    result = _download_pdf_from_urls(screener_urls, save_path)
    if result:
        return result

    # Strategy 2: BSE/NSE direct known URLs (Fallback 1)
    direct_urls = _try_bse_nse_direct_download(symbol, financial_year)
    result = _download_pdf_from_urls(direct_urls, save_path)
    if result:
        return result

    # Strategy 3: Web search (Fallback 2)
    search_queries = [
        f"{symbol} annual report {financial_year} filetype:pdf site:bseindia.com OR site:nseindia.com OR site:moneycontrol.com",
        f"{symbol} annual report {financial_year} filetype:pdf",
        f"{symbol} annual report FY{financial_year} PDF download",
    ]
    all_urls: List[str] = []
    for q in search_queries:
        logger.info(f"Searching: {q}")
        found = search_internet_for_pdf_links(q)
        all_urls.extend(found)
        if all_urls:
            break  # Stop after first successful search result set

    if not all_urls:
        logger.warning(f"No PDF links found for {symbol} FY{financial_year} via any search strategy.")
        return None

    logger.info(f"Found {len(all_urls)} total PDF candidate(s). Starting download attempts...")
    result = _download_pdf_from_urls(all_urls, save_path)
    if not result:
        logger.error(f"All download attempts failed for {symbol} FY{financial_year}.")
    return result


def fetch_corporate_document(
    symbol: str,
    document_type: str,
    financial_year: int,
    quarter: Optional[str] = None
) -> Optional[str]:
    """
    Download a specific corporate document (quarterly result, concall, presentation)
    for a stock. Uses Screener.in first, then falls back to web searches.
    """
    q_str = f" {quarter}" if quarter else ""
    logger.info(f"Fetching {document_type}{q_str} FY{financial_year} for {symbol}...")

    doc_dir = os.path.join(settings.DATA_DIR, "documents")
    os.makedirs(doc_dir, exist_ok=True)
    q_suffix = f"_{quarter}" if quarter else ""
    file_name = f"{symbol.upper()}_{financial_year}{q_suffix}_{document_type}.pdf"
    save_path = os.path.join(doc_dir, file_name)

    # Shortcut: If valid PDF already exists, reuse it
    if os.path.exists(save_path) and os.path.getsize(save_path) > 10000:
        try:
            with open(save_path, "rb") as f:
                magic = f.read(4)
                if magic == b"%PDF":
                    logger.info(f"PDF already exists on disk: {save_path}. Skipping download.")
                    return save_path
        except Exception:
            pass

    # Strategy 1: Screener.in scraper (Primary)
    screener_urls = _get_screener_pdf_links(symbol, document_type, financial_year, quarter)
    result = _download_pdf_from_urls(screener_urls, save_path)
    if result:
        return result

    # Strategy 2: Web search queries (Fallback)
    # Build query variations based on document type
    if document_type == "quarterly_result":
        queries = [
            f"{symbol} quarterly results{q_str} {financial_year} filetype:pdf site:bseindia.com",
            f"{symbol} quarterly financial results{q_str} {financial_year} filetype:pdf",
            f"{symbol} results{q_str} {financial_year} filetype:pdf",
            f"{symbol} financial statements{q_str} {financial_year} PDF",
        ]
    elif document_type == "concall":
        queries = [
            f"{symbol} concall transcript{q_str} {financial_year} filetype:pdf",
            f"{symbol} earnings call transcript{q_str} {financial_year} filetype:pdf",
            f"{symbol} investor call{q_str} {financial_year} transcript PDF",
        ]
    elif document_type == "presentation":
        queries = [
            f"{symbol} investor presentation{q_str} {financial_year} filetype:pdf",
            f"{symbol} analyst presentation{q_str} {financial_year} filetype:pdf",
            f"{symbol} earnings presentation{q_str} {financial_year} PDF",
        ]
    else:
        queries = [
            f"{symbol} annual report {financial_year} filetype:pdf",
        ]

    # Collect PDF URLs from the first search query that yields results
    all_urls: List[str] = []
    for q in queries:
        logger.info(f"Searching: {q}")
        found = search_internet_for_pdf_links(q)
        all_urls.extend(found)
        if all_urls:
            break

    if not all_urls:
        logger.warning(f"No PDF links found for {symbol} {document_type}{q_str} FY{financial_year}.")
        return None

    logger.info(f"Found {len(all_urls)} PDF candidate(s). Starting downloads...")
    result = _download_pdf_from_urls(all_urls, save_path)
    if not result:
        logger.error(f"All download attempts failed for {symbol} {document_type}{q_str} FY{financial_year}.")
    return result
